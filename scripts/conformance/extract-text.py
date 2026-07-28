#!/usr/bin/env python3
"""Ground-truth text extractor for the reader-fidelity harness.

Prints every non-empty paragraph and table-cell text of a DOCX (body,
headers, footers), one per line, whitespace-normalized — extracted by
python-docx, i.e. by an implementation independent of php-docx.

Cached complex-field results (runs between fldChar separate and end,
e.g. the frozen page number of a PAGE field) are excluded: they are a
render-time artifact, not authored text — php-docx models the field
itself, not its stale cache.

Usage: python3 extract-text.py <file.docx>
"""

import re
import sys

import docx
from docx.oxml.ns import qn


def norm(s: str) -> str:
    return re.sub(r"\s+", " ", s).strip()


def emit(text: str) -> None:
    n = norm(text)
    if n:
        print(n)


def _inside_fld_simple(node) -> bool:
    parent = node.getparent()
    while parent is not None:
        if parent.tag == qn("w:fldSimple"):
            return True
        parent = parent.getparent()
    return False


def para_text(p) -> str:
    """Paragraph text with cached field values skipped (fldChar + fldSimple)."""
    out = []
    state = None  # None | "instr" | "value"
    for node in p._p.iter():
        tag = node.tag
        if tag == qn("w:fldChar"):
            t = node.get(qn("w:fldCharType"))
            if t == "begin":
                state = "instr"
            elif t == "separate":
                state = "value"
            elif t == "end":
                state = None
        elif tag == qn("w:t") and state is None and not _inside_fld_simple(node):
            out.append(node.text or "")
        elif tag in (qn("w:tab"), qn("w:br")) and state is None:
            out.append(" ")
    return "".join(out)


def cell_text(cell) -> str:
    return " ".join(para_text(p) for p in cell.paragraphs)


d = docx.Document(sys.argv[1])

for p in d.paragraphs:
    emit(para_text(p))
for t in d.tables:
    for row in t.rows:
        for cell in row.cells:
            emit(cell_text(cell))
for section in d.sections:
    for container in (section.header, section.footer):
        for p in container.paragraphs:
            emit(para_text(p))
