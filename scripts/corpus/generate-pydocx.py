#!/usr/bin/env python3
"""Corpus generator: DOCX produced by python-docx — an independent
producer with its own OOXML emission quirks.

Usage: python3 scripts/corpus/generate-pydocx.py [out-dir]
  default out-dir: build/corpus
"""

import sys
from pathlib import Path

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH

out_dir = Path(sys.argv[1] if len(sys.argv) > 1 else
               Path(__file__).resolve().parents[2] / "build" / "corpus")
out_dir.mkdir(parents=True, exist_ok=True)

d = docx.Document()
d.add_heading("python-docx document", level=1)
p = d.add_paragraph("Plain paragraph. ")
p.add_run("Bold run").bold = True
p.add_run(" and ")
p.add_run("italic run").italic = True
p.add_run(". Кириллица и ünïcode.")

j = d.add_paragraph("Justified paragraph produced by python-docx with "
                    "enough words to matter.")
j.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

d.add_heading("Table", level=2)
t = d.add_table(rows=2, cols=3)
t.style = "Table Grid"
for ci, text in enumerate(["Alpha", "Beta", "Gamma"]):
    t.rows[0].cells[ci].text = text
merged = t.rows[1].cells[0].merge(t.rows[1].cells[1])
merged.text = "merged horizontally"
t.rows[1].cells[2].text = "C2"

d.add_heading("Lists", level=2)
d.add_paragraph("bullet one", style="List Bullet")
d.add_paragraph("bullet nested", style="List Bullet 2")
d.add_paragraph("numbered one", style="List Number")
d.add_paragraph("numbered two", style="List Number")

d.save(out_dir / "pydocx-features.docx")
print(f"generated pydocx-features.docx in {out_dir}")
